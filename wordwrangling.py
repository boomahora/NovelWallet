import random
import win32com.client as win32
from docx import Document
import PyPDF2
import re
from fuzzywuzzy import process
from docx.shared import Pt


def modify_original_word(doc_path, modified_doc_path, words):
    """
    Modify the original Word document by replacing selected words with identifiers, and
    convert the modified document to PDF. The function returns chapter ranges and the PDF path.

    :param doc_path: Path to the original Word document.
    :param modified_doc_path: Path to save the modified Word document.
    :param words: List of words to replace.
    :return: A tuple containing chapter ranges and the path to the generated PDF.
    """
    document = Document(doc_path)
    chapter_starts_paragraphs = [i for i, p in enumerate(document.paragraphs) if p.style.name == 'CSP - Chapter Title']

    for idx, chapter_start in enumerate(chapter_starts_paragraphs):
        end_idx = chapter_starts_paragraphs[idx+1] if idx < len(chapter_starts_paragraphs) - 1 else len(document.paragraphs)
        replace_words_with_identifier(chapter_start, end_idx, words, document)

    document.save(modified_doc_path)

    # Convert the Word document to PDF
    word = win32.Dispatch("Word.Application")
    doc = word.Documents.Open(modified_doc_path)

    # Embed fonts
    doc.EmbedTrueTypeFonts = True
    doc.SaveSubsetFonts = True

    pdf_path = modified_doc_path.replace('.docx', '.pdf')
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 represents the file format for PDFs in Word's API

    doc.Close()
    word.Quit()
    # Extract chapter page ranges
    chapter_dict = read_pdf_pages(pdf_path)
    chapter_dict = fix_chapter_dict(chapter_dict)
    
    chapter_ranges = {}
    for chapter, pages in chapter_dict.items():
        chapter_ranges[chapter] = list(range(min(pages), max(pages) + 1))

    return chapter_ranges, pdf_path



def replace_words_with_identifier(start, end, words, document):
    """
    In the provided range within the document, replace random words with the provided list of words, 
    wrapping them with identifiers.

    :param start: Start index of the paragraph range.
    :param end: End index of the paragraph range.
    :param words: List of words to insert.
    :param document: The document object.
    """
    word_list = words.copy()
    random.shuffle(word_list)
    
    replaceable_paragraphs = document.paragraphs[start+1:end]
    replaceable_texts = [p.text for p in replaceable_paragraphs]
    flat_replaceable_texts = ' '.join(replaceable_texts).split()
    
    replaced_indices = random.sample(range(len(flat_replaceable_texts)), len(word_list))
    for i, word_to_replace in enumerate(word_list):
        current_word = flat_replaceable_texts[replaced_indices[i]]
        punctuation_end = ''
        punctuation_start = ''

        if current_word[-1] in [".", ",", ":", ";", "!", "?","“"]:

            punctuation_end = current_word[-1]
        
        if current_word[0] == '"':

            punctuation_start = '“'

        flat_replaceable_texts[replaced_indices[i]] = f"{punctuation_start}|{word_to_replace}|{punctuation_end}"
    
    word_idx = 0
    for paragraph in replaceable_paragraphs:
        new_paragraph = []
        for _ in paragraph.text.split():
            new_paragraph.append(flat_replaceable_texts[word_idx])
            word_idx += 1
        paragraph.text = ' '.join(new_paragraph)


def extract_words_from_pdf(pdf_path):
    """
    Extract words with identifiers from the provided PDF.

    :param pdf_path: Path to the PDF file.
    :return: A list of tuples containing extracted word, page number, chapter, and page index.
    """

    chapter_dict = fix_chapter_dict(read_pdf_pages(pdf_path))
    
    def get_chapter_for_page(page_num, chapter_dict):
        for chapter, pages in chapter_dict.items():
            if page_num in pages:
                return chapter
        return None

    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        extracted_words = []

        # Regular expression to find our words, allowing for spaces, linebreaks, etc.
        word_pattern = re.compile(r'\|\s*([^|]+)\s*\|')
        page_num_pattern = re.compile(r'^(?:Chapter \d+\s+)?(\d+)', re.MULTILINE)

        last_page = 0
        last_chapter = 0

        for page_num in range(len(reader.pages)):
            
            page = reader.pages[page_num]
            page_text = page.extract_text().replace('\n', ' ')

            if not page_text.split():
                continue
            
            first_word = page_text.split()[0]

            if first_word.lower() == "chapter":
                actual_page_num = int(page_text.split()[2])
            elif first_word.isdigit():
                actual_page_num = int(first_word)
            else:
                actual_page_num = page_num + 1


            # Compare with last_page and adjust if necessary
            if actual_page_num <= last_page:
                actual_page_num = last_page + 1

            # Determine the chapter for this page
            current_chapter = get_chapter_for_page(actual_page_num, chapter_dict)

            # Compare with last_chapter and adjust if necessary
            if current_chapter and current_chapter <= last_chapter:
                current_chapter = last_chapter

            # Update last_page and last_chapter for the next iteration
            last_page = actual_page_num
            last_chapter = current_chapter if current_chapter else last_chapter

            # Calculate the pageindex
            page_position_in_chapter = chapter_dict[current_chapter].index(actual_page_num) + 1
            pageindex = page_position_in_chapter 

            # Find all matches on the page
            for match in word_pattern.finditer(page_text):
                word = match.group(1).replace(" ", "").strip("|")
                extracted_words.append((word, actual_page_num, current_chapter, pageindex))

    return extracted_words


def map_extracted_words(extracted_words, original_words, old_chapter_ranges, new_chapter_ranges):
    """
    Map extracted words to their original form and return the new page number, chapter, and page index.

    :param extracted_words: List of words extracted from the PDF.
    :param original_words: List of original words before modification.
    :param old_chapter_ranges: Chapter ranges from the original document.
    :param new_chapter_ranges: Chapter ranges from the modified document.
    :return: A list of tuples containing mapped word, new page number, and chapter.
    """
    mapped_words = []

    for word, page_num, chapter, pageindex in extracted_words:
                
        if not word or not any(c.isalnum() for c in word):
            continue

        best_match, score = process.extractOne(word, original_words)

        if score > 80:  # You can adjust the threshold as needed
            new_page_num = get_new_page_num(pageindex, old_chapter_ranges, new_chapter_ranges, chapter)
            mapped_words.append((best_match, new_page_num, chapter))

    # Error handling
    expected_num_words = len(old_chapter_ranges) * len(original_words)
    if len(extracted_words) != expected_num_words:
        print(f'''Error: Expected {expected_num_words} words, but extracted {len(extracted_words)} words.\n 
              This can happen when the last page has no header, but still a footer. Does not affect the program.''')

    return mapped_words

def remove_identifiers_from_docx(doc_path, save_path):
    """
    Remove identifiers from the Word document.

    :param doc_path: Path to the Word document with identifiers.
    :param save_path: Path to save the cleaned Word document.
    """
    # Load the Word document using python-docx
    doc = Document(doc_path)
    
    # Iterate through each paragraph and replace ##
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("|", "")
    
    # Save the modified document
    doc.save(save_path)

def set_font_for_docx(doc_path, save_as_pdf=True):
    """
    Set the font for the Word document. Optionally, convert the Word document to PDF.

    :param doc_path: Path to the Word document.
    :param save_as_pdf: Boolean indicating whether to convert the document to PDF.
    """
    # Load the Word document using python-docx
    doc = Document(doc_path)

    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        # Skip headings based on style name (assuming standard heading styles; adjust as needed)
        if paragraph.style.name.startswith('Heading'):
            continue
        
        # Iterate through each run in the paragraph
        for run in paragraph.runs:
            # Set font and size
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    # Save the modified document
    doc.save(doc_path)

    # Convert to PDF if required
    if save_as_pdf:
        word = win32.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_path)

        # Define the path for the PDF file
        pdf_path = doc_path.replace('.docx', '.pdf')

        # Convert the Word document to PDF
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 represents the file format for PDFs in Word's API

        doc.Close()
        word.Quit()



def read_pdf_pages(pdf_path):
    with open(pdf_path, "rb") as file:
        pdf_reader = PyPDF2.PdfReader(file)
        
        chapter_dict = {}
        current_chapter = 0  

        # Iterate over all pages
        for page_num in range(len(pdf_reader.pages)):
            page_text = pdf_reader.pages[page_num].extract_text().strip()

            if len(page_text) == 0:
                continue
            else:
                first_word = page_text.split()[0]
                
                # If the page starts with the word "Chapter", it's a continuation of the current chapter
                if first_word.lower() == "chapter":
                    actual_page_num = int(page_text.split()[2])
                    chapter_dict.setdefault(current_chapter, []).append(actual_page_num)

                # If the page starts with a number (and doesn't have the word "Chapter"), it's a new chapter
                elif first_word.isdigit():
                    current_chapter += 1  # Increment the chapter number
                    actual_page_num = int(first_word)
                    chapter_dict[current_chapter] = [actual_page_num]

        return chapter_dict



def get_new_page_num(pageindex, old_chapter_ranges, new_chapter_ranges, chapter):
    """
    Get the new page number by mapping the pageindex from the old chapter ranges to the new chapter ranges.

    :param pageindex: Index of the page in the chapter.
    :param old_chapter_ranges: Chapter ranges from the original document.
    :param new_chapter_ranges: Chapter ranges from the modified document.
    :param chapter: Chapter number.
    :return: New page number.
    """

    old_pages = old_chapter_ranges[chapter]
    new_pages = new_chapter_ranges[chapter]

    # Get the old page number using the pageindex
    old_page_num = old_pages[pageindex - 1]

    # Find the index of the old page number in the old_pages list
    old_index = old_pages.index(old_page_num)

    # Use the same index to get the corresponding new page number from the new_pages list
    # If the index is out of bounds, get the last page number from the new_pages list
    new_page_num = new_pages[old_index] if old_index < len(new_pages) else new_pages[-1]

    return new_page_num



def fix_chapter_dict(chapter_dict):
    # Sort the chapter_dict by keys (chapter numbers)
    sorted_chapters = sorted(chapter_dict.keys())

    # Initialize the previous chapter and page values for comparison
    prev_chapter = 0
    prev_page = 0

    fixed_chapter_dict = {}

    for chapter in sorted_chapters:
        if chapter <= prev_chapter:
            chapter = prev_chapter + 1  # Ensure the chapter number is always increasing

        # Sort the pages for the current chapter
        sorted_pages = sorted(chapter_dict[chapter])

        fixed_pages = []
        for page in sorted_pages:
            if page <= prev_page:
                page = prev_page + 1  # Ensure the page number is always increasing
            fixed_pages.append(page)
            prev_page = page

        fixed_chapter_dict[chapter] = fixed_pages
        prev_chapter = chapter

    return fixed_chapter_dict
